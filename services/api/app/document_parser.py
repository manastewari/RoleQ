import re
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import HTTPException
from pypdf import PdfReader


MAX_DOCUMENT_CHARS = 120_000


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_DOCUMENT_CHARS]


def parse_document(path: Path, filename: str) -> str:
    return parse_document_bytes(path.read_bytes(), filename)


def parse_document_bytes(content: bytes, filename: str) -> str:
    extension = Path(filename).suffix.lower()
    if extension == ".doc":
        raise HTTPException(
            status_code=415,
            detail="Legacy .doc files are not supported. Convert the file to DOCX or PDF and upload it again.",
        )

    try:
        if extension == ".pdf":
            reader = PdfReader(BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            text = clean_text(text)
            if len(text) < 40:
                raise HTTPException(
                    status_code=422,
                    detail="This PDF appears to be scanned or image-only. OCR is not included in this prototype.",
                )
            return text

        if extension == ".docx":
            document = Document(BytesIO(content))
            blocks = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    blocks.append(" | ".join(cell.text for cell in row.cells))
            text = clean_text("\n".join(blocks))
            if len(text) < 20:
                raise HTTPException(status_code=422, detail="The DOCX did not contain enough readable text.")
            return text

        if extension == ".txt":
            try:
                return clean_text(content.decode("utf-8"))
            except UnicodeDecodeError:
                return clean_text(content.decode("latin-1", errors="replace"))
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not parse {extension or 'document'}: {exc}") from exc

    raise HTTPException(status_code=415, detail="Only PDF, DOCX, and TXT files are supported.")
